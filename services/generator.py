"""
Resume & cover letter generation — tailored to specific job postings using LLM.
"""

import json
import os
import re
import sqlite3
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
from services.llm import llm_chat
from services.resumes import load_candidate_profile
from services.settings import load_settings

GENERATED_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "data", "generated")


def _md_to_docx(markdown: str, output_path: str, doc_type: str = "resume"):
    """Convert markdown text to a clean .docx file."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Tighter paragraph spacing
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(0)

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    lines = markdown.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Horizontal rule → thin line (skip it, just adds spacing)
        if stripped in ("---", "***", "___"):
            i += 1
            continue

        # H1: # Header
        if stripped.startswith("# ") and not stripped.startswith("## "):
            text = stripped[2:].strip()
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)  # strip bold markers
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # H2: ## Section Header
        if stripped.startswith("## "):
            text = stripped[3:].strip()
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            p = doc.add_paragraph()
            run = p.add_run(text.upper())
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x1a, 0x56, 0x8a)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            # Add a bottom border
            from docx.oxml.ns import qn
            pPr = p._p.get_or_add_pPr()
            pBdr = pPr.makeelement(qn("w:pBdr"), {})
            bottom = pBdr.makeelement(qn("w:bottom"), {
                qn("w:val"): "single",
                qn("w:sz"): "4",
                qn("w:space"): "1",
                qn("w:color"): "1a568a",
            })
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # H3: ### Sub-header
        if stripped.startswith("### "):
            text = stripped[4:].strip()
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(11)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Bullet point: - text or * text
        if re.match(r"^[-*]\s", stripped):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(p, text)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Inches(0.25)
            i += 1
            continue

        # Regular paragraph (handle bold/italic inline)
        p = doc.add_paragraph()
        _add_formatted_text(p, stripped)
        # Center lines that look like contact info (short, contain | or email)
        if len(stripped) < 120 and ("|" in stripped or "@" in stripped):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(1)
        i += 1

    doc.save(output_path)


def _add_formatted_text(paragraph, text):
    """Add text to a paragraph, handling **bold** and *italic* markers."""
    # Split on bold markers first
    parts = re.split(r"(\*\*[^*]+?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif "*" in part:
            # Handle italic within non-bold segments
            sub_parts = re.split(r"(\*[^*]+?\*)", part)
            for sp in sub_parts:
                if sp.startswith("*") and sp.endswith("*"):
                    run = paragraph.add_run(sp[1:-1])
                    run.italic = True
                else:
                    if sp:
                        paragraph.add_run(sp)
        else:
            if part:
                paragraph.add_run(part)


def _md_to_pdf(markdown: str, output_path: str):
    """Convert markdown text to a clean PDF file."""
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    pdf.set_margins(20, 15, 20)

    lines = markdown.split("\n")
    for line in lines:
        stripped = line.strip()

        # Skip empty lines — add small spacing
        if not stripped:
            pdf.ln(3)
            continue

        # Horizontal rule
        if stripped in ("---", "***", "___"):
            y = pdf.get_y()
            pdf.set_draw_color(180, 180, 180)
            pdf.line(20, y, 190, y)
            pdf.ln(3)
            continue

        # Strip bold/italic markers for PDF (fpdf2 doesn't do inline mixed easily)
        clean = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
        clean = re.sub(r"\*(.+?)\*", r"\1", clean)

        # H1
        if stripped.startswith("# ") and not stripped.startswith("## "):
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped[2:].strip())
            pdf.set_font("Helvetica", "B", 16)
            pdf.set_text_color(26, 26, 26)
            pdf.cell(0, 8, text, align="C", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1)
            continue

        # H2
        if stripped.startswith("## "):
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped[3:].strip())
            pdf.ln(3)
            pdf.set_font("Helvetica", "B", 12)
            pdf.set_text_color(26, 86, 138)
            pdf.cell(0, 7, text.upper(), new_x="LMARGIN", new_y="NEXT")
            # Underline
            y = pdf.get_y()
            pdf.set_draw_color(26, 86, 138)
            pdf.line(20, y, 190, y)
            pdf.ln(2)
            continue

        # H3
        if stripped.startswith("### "):
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped[4:].strip())
            pdf.ln(2)
            pdf.set_font("Helvetica", "B", 11)
            pdf.set_text_color(51, 51, 51)
            pdf.cell(0, 6, text, new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1)
            continue

        # Bullet point
        if re.match(r"^[-*]\s", stripped):
            text = clean[2:].strip()
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(51, 51, 51)
            x = pdf.get_x()
            pdf.cell(8, 5, chr(8226), new_x="END")  # bullet char
            pdf.multi_cell(0, 5, text, new_x="LMARGIN", new_y="NEXT")
            continue

        # Regular paragraph
        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(51, 51, 51)
        # Center contact-like lines
        if len(clean) < 120 and ("|" in clean or "@" in clean):
            pdf.cell(0, 5, clean, align="C", new_x="LMARGIN", new_y="NEXT")
        else:
            pdf.multi_cell(0, 5, clean, new_x="LMARGIN", new_y="NEXT")

    pdf.output(output_path)


def _pick_best_resume(job: dict) -> tuple[str, str]:
    """Pick the best-matching uploaded resume for a job.

    Compares each resume's best_for tags against the job title/description.
    Returns (content_text, original_name) of the best match.
    Falls back to the longest resume if no tag matches.
    """
    db_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "data", "jobs.db")
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    rows = conn.execute(
        "SELECT original_name, content_text, best_for FROM resumes WHERE content_text IS NOT NULL"
    ).fetchall()
    conn.close()

    if not rows:
        return "", ""

    job_text = f"{job.get('title', '')} {job.get('description', '')}".lower()

    best_score = -1
    best_row = None
    longest_row = None
    longest_len = 0

    for row in rows:
        text = row["content_text"] or ""

        # Track longest as fallback
        if len(text) > longest_len:
            longest_len = len(text)
            longest_row = row

        # Score by best_for tag overlap
        best_for = row["best_for"]
        if best_for:
            try:
                tags = json.loads(best_for) if isinstance(best_for, str) else best_for
            except (json.JSONDecodeError, TypeError):
                tags = []

            score = 0
            for tag in tags:
                # Check if any word from the tag appears in job text
                tag_words = tag.lower().split()
                for word in tag_words:
                    if len(word) > 3 and word in job_text:
                        score += 1
            if score > best_score:
                best_score = score
                best_row = row

    chosen = best_row if best_score > 0 else longest_row
    if chosen:
        return chosen["content_text"], chosen["original_name"]
    return "", ""


def generate_resume(job: dict, settings: dict = None) -> dict:
    """Generate a tailored resume for a specific job posting.

    Returns {ok, markdown, path, resume_source} or {error}.
    """
    if not settings:
        settings = load_settings()

    profile = load_candidate_profile()
    if not profile:
        return {"error": "No candidate profile. Analyze resumes first."}

    resume_text, resume_name = _pick_best_resume(job)
    if not resume_text:
        return {"error": "No resumes uploaded. Upload resumes in Arsenal first."}

    job_info = (
        f"Title: {job.get('title', 'Unknown')}\n"
        f"Company: {job.get('company', 'Unknown')}\n"
        f"Location: {job.get('location', 'Unknown')}\n"
        f"Industry: {job.get('industry', 'Unknown')}\n"
        f"Salary: {job.get('salary_text', 'Not listed')}\n"
        f"Description:\n{(job.get('description', '') or '')[:4000]}"
    )

    # Build work history timeline for the prompt
    work_timeline = ""
    for wh in profile.get("work_history", []):
        start = wh.get("start_year", "")
        end = wh.get("end_year", "")
        date_range = f"{start}-{end}" if start else wh.get("duration", "")
        location = wh.get("location", "")
        loc_str = f" ({location})" if location else ""
        work_timeline += f"- {wh['title']} @ {wh['company']}{loc_str} [{date_range}]\n"

    prompt = f"""You are an expert resume writer. Create a tailored resume for this specific job posting.

CANDIDATE PROFILE:
Name: {profile.get('name', 'Unknown')}
Headline: {profile.get('headline', '')}
Experience: {profile.get('experience_years', 0)}+ years
Core Strengths: {', '.join(profile.get('core_strengths', []))}
Key Skills: {', '.join(profile.get('all_skills', [])[:25])}
Unique Value: {profile.get('unique_value', '')}

WORK TIMELINE (include dates on resume — do NOT omit any current positions):
{work_timeline}

SOURCE RESUME (use this as the base — reorder, emphasize, and tailor):
{resume_text[:5000]}

TARGET JOB:
{job_info}

INSTRUCTIONS:
- Tailor this resume specifically for the job above
- Lead with the most relevant experience and skills for THIS role
- Use keywords and phrases from the job description naturally
- Keep it to 1-2 pages worth of content
- Use clean markdown formatting with clear sections
- Include: Contact header, Professional Summary (tailored), Key Skills (relevant ones first), Professional Experience (reordered by relevance), Education
- Do NOT fabricate experience or skills — only reorganize and emphasize what's real
- Make the professional summary directly address what this employer is looking for

Output the resume in clean markdown. No commentary before or after — just the resume."""

    result = llm_chat(
        [{"role": "user", "content": prompt}],
        settings,
    )

    # Save markdown
    os.makedirs(GENERATED_DIR, exist_ok=True)
    md_filename = f"{job['id']}_resume.md"
    md_path = os.path.join(GENERATED_DIR, md_filename)
    with open(md_path, "w") as f:
        f.write(result)

    # Convert to DOCX and PDF
    docx_path = os.path.join(GENERATED_DIR, f"{job['id']}_resume.docx")
    pdf_path = os.path.join(GENERATED_DIR, f"{job['id']}_resume.pdf")
    _md_to_docx(result, docx_path, doc_type="resume")
    try:
        _md_to_pdf(result, pdf_path)
    except Exception:
        pass  # PDF is optional — Unicode font issues shouldn't block generation

    # Update job record (store base path without extension)
    base_path = os.path.join(GENERATED_DIR, f"{job['id']}_resume")
    db_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "data", "jobs.db")
    conn = sqlite3.connect(db_path)
    conn.execute(
        "UPDATE jobs SET resume_path = ?, updated_at = datetime('now') WHERE id = ?",
        (base_path, job["id"]),
    )
    conn.commit()
    conn.close()

    return {
        "ok": True,
        "markdown": result,
        "path": base_path,
        "resume_source": resume_name,
    }


def generate_cover_letter(job: dict, settings: dict = None) -> dict:
    """Generate a tailored cover letter for a specific job posting.

    Returns {ok, markdown, path} or {error}.
    """
    if not settings:
        settings = load_settings()

    profile = load_candidate_profile()
    if not profile:
        return {"error": "No candidate profile. Analyze resumes first."}

    elevator_pitch = settings.get("candidate_elevator_pitch", "")

    job_info = (
        f"Title: {job.get('title', 'Unknown')}\n"
        f"Company: {job.get('company', 'Unknown')}\n"
        f"Location: {job.get('location', 'Unknown')}\n"
        f"Description:\n{(job.get('description', '') or '')[:4000]}"
    )

    work_history = ""
    for wh in profile.get("work_history", [])[:5]:
        highlights = "; ".join(wh.get("highlights", [])[:2])
        start = wh.get("start_year", "")
        end = wh.get("end_year", "")
        date_range = f"{start}-{end}" if start else wh.get("duration", "")
        location = wh.get("location", "")
        loc_str = f", {location}" if location else ""
        work_history += f"- {wh['title']} @ {wh['company']}{loc_str} ({date_range}): {highlights}\n"

    prompt = f"""You are an expert career coach and cover letter writer. Write a professional, compelling cover letter for this specific job.

CANDIDATE:
Name: {profile.get('name', 'Unknown')}
Headline: {profile.get('headline', '')}
Experience: {profile.get('experience_years', 0)}+ years
Core Strengths: {', '.join(profile.get('core_strengths', []))}
Key Skills: {', '.join(profile.get('all_skills', [])[:20])}
Unique Value: {profile.get('unique_value', '')}
{f'Elevator Pitch: {elevator_pitch}' if elevator_pitch else ''}

RELEVANT WORK HISTORY:
{work_history}

TARGET JOB:
{job_info}

INSTRUCTIONS:
- Write a professional cover letter (3-4 paragraphs)
- Opening: Hook that shows you understand what this company/role needs
- Body: Connect 2-3 specific experiences to what the job requires — be concrete, not generic
- Closing: Confident but not arrogant, express genuine interest, call to action
- Tone: Professional, direct, personable — not stuffy corporate-speak
- Do NOT use cliches like "I'm writing to express my interest" or "I believe I would be an asset"
- Do NOT fabricate anything — only reference real experience from the profile
- Keep it under 400 words
- Use markdown formatting

Output only the cover letter. No commentary before or after."""

    result = llm_chat(
        [{"role": "user", "content": prompt}],
        settings,
    )

    # Save markdown
    os.makedirs(GENERATED_DIR, exist_ok=True)
    md_filename = f"{job['id']}_cover.md"
    md_path = os.path.join(GENERATED_DIR, md_filename)
    with open(md_path, "w") as f:
        f.write(result)

    # Convert to DOCX and PDF
    docx_path = os.path.join(GENERATED_DIR, f"{job['id']}_cover.docx")
    pdf_path = os.path.join(GENERATED_DIR, f"{job['id']}_cover.pdf")
    _md_to_docx(result, docx_path, doc_type="cover")
    try:
        _md_to_pdf(result, pdf_path)
    except Exception:
        pass  # PDF is optional — Unicode font issues shouldn't block generation

    # Update job record (store base path without extension)
    base_path = os.path.join(GENERATED_DIR, f"{job['id']}_cover")
    db_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "data", "jobs.db")
    conn = sqlite3.connect(db_path)
    conn.execute(
        "UPDATE jobs SET cover_letter_path = ?, updated_at = datetime('now') WHERE id = ?",
        (base_path, job["id"]),
    )
    conn.commit()
    conn.close()

    return {
        "ok": True,
        "markdown": result,
        "path": base_path,
    }


def revise_document(job: dict, doc_type: str, instruction: str,
                    settings: dict = None) -> dict:
    """Revise a previously generated resume or cover letter based on user feedback.

    doc_type: "resume" or "cover"
    instruction: user's revision note (e.g. "tone down the HIWC section")
    Returns {ok, markdown, path} or {error}.
    """
    if not settings:
        settings = load_settings()

    if doc_type not in ("resume", "cover"):
        return {"error": "Invalid doc_type. Use 'resume' or 'cover'."}

    # Find the current markdown on disk
    field = "resume_path" if doc_type == "resume" else "cover_letter_path"
    base_path = job.get(field, "")
    if not base_path:
        return {"error": f"No {doc_type} generated yet. Generate one first."}

    if base_path.endswith((".docx", ".pdf", ".md")):
        base_path = os.path.splitext(base_path)[0]

    md_path = f"{base_path}.md"
    if not os.path.exists(md_path):
        return {"error": f"Source file not found at {md_path}. Try regenerating."}

    with open(md_path) as f:
        current_doc = f.read()

    # Build the job context for the LLM
    job_info = (
        f"Title: {job.get('title', 'Unknown')}\n"
        f"Company: {job.get('company', 'Unknown')}\n"
        f"Location: {job.get('location', 'Unknown')}\n"
    )

    doc_label = "resume" if doc_type == "resume" else "cover letter"

    prompt = f"""You are an expert {doc_label} editor. The user has a {doc_label} that was previously generated for a specific job posting. They want you to revise it based on their feedback.

TARGET JOB:
{job_info}

CURRENT {doc_label.upper()}:
{current_doc}

USER'S REVISION REQUEST:
{instruction}

INSTRUCTIONS:
- Apply the user's requested changes to the {doc_label}
- Keep all other content and formatting intact unless the change requires restructuring
- Maintain the same markdown formatting style
- Do NOT fabricate new experience or skills — only adjust what's already there
- Do NOT add commentary before or after — output ONLY the revised {doc_label}"""

    result = llm_chat(
        [{"role": "user", "content": prompt}],
        settings,
    )

    # Overwrite the markdown file
    with open(md_path, "w") as f:
        f.write(result)

    # Regenerate DOCX and PDF from the revised markdown
    docx_path = f"{base_path}.docx"
    pdf_path = f"{base_path}.pdf"
    _md_to_docx(result, docx_path, doc_type=doc_type)
    try:
        _md_to_pdf(result, pdf_path)
    except Exception:
        pass  # PDF is optional

    return {
        "ok": True,
        "markdown": result,
        "path": base_path,
    }
